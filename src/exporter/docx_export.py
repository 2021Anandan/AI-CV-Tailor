import io
from docx import Document

def create_docx(text_content):
    doc = Document()
    doc.add_heading("Optimized Resume", level=1)
    
    # Split text by lines and add paragraphs cleanly
    for line in text_content.split("\n"):
        if line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph()
            
    # Save to a byte stream
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream